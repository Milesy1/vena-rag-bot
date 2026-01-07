"""Create POC Findings Report"""
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

doc = Document()

# Title
doc.add_heading('Vena RAG Bot - POC Findings Report', 0)
doc.add_paragraph(f'Date: {datetime.now().strftime("%d %B %Y")}')
doc.add_paragraph('Author: Miles Sherwood')
doc.add_paragraph('Status: POC Complete')

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph('A proof of concept RAG (Retrieval Augmented Generation) bot was successfully built and deployed to demonstrate AI-powered technical support for the Vena platform.')
doc.add_paragraph('Key Findings:')
findings = [
    'The system successfully retrieves relevant information from a knowledge base of 27 Vena documents',
    'Semantic search correctly matches user questions to relevant content even without exact keyword matches',
    'Source citations provide transparency and traceability for all answers',
    'The solution can be deployed to a shareable URL for stakeholder demonstration',
    'Technical constraints were identified and solutions documented for production deployment'
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Recommendation: GO - Proceed with Phase 1 implementation')

# 2. Test Results
doc.add_heading('2. Test Results', level=1)

doc.add_heading('2.1 Quantitative Metrics', level=2)
metrics = [
    ('Documents Ingested', '27'),
    ('Document Types', 'Explainers, How-To Guides, Reference, Troubleshooting'),
    ('Average Response Time', '3-5 seconds'),
    ('Knowledge Base Build Time', '60-90 seconds'),
    ('Deployment Platform', 'Streamlit Community Cloud'),
    ('LLM Model', 'GPT-4o'),
    ('Embedding Model', 'text-embedding-3-small'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Value'
for metric, value in metrics:
    row = table.add_row().cells
    row[0].text = metric
    row[1].text = value

doc.add_heading('2.2 Test Queries', level=2)
doc.add_paragraph('Sample questions tested successfully:')
queries = [
    'What is the difference between alternate hierarchies and calculated members? - PASS',
    'How do I write a basic VenaQL query? - PASS',
    'My data is not loading, what should I check? - PASS',
    'How do I handle FX conversion? - PASS',
]
for q in queries:
    doc.add_paragraph(q, style='List Bullet')

# 3. Technical Learnings
doc.add_heading('3. Technical Learnings', level=1)

doc.add_heading('3.1 What Worked Well', level=2)
worked = [
    'LangChain framework provided rapid development of RAG pipeline',
    'ChromaDB vector database enabled effective semantic search',
    'OpenAI embeddings produced quality document matching',
    'Streamlit provided quick path to deployable web interface',
    'PDF to markdown conversion preserved document structure',
]
for w in worked:
    doc.add_paragraph(w, style='List Bullet')

doc.add_heading('3.2 Challenges Encountered and Resolved', level=2)
challenges = [
    'Streamlit Cloud read-only filesystem - Resolved with in-memory vector storage',
    'API key configuration for cloud - Resolved using Streamlit secrets',
    'Knowledge base rebuild required after updates - Expected for in-memory storage',
]
for c in challenges:
    doc.add_paragraph(c, style='List Bullet')

# 4. Knowledge Base Quality
doc.add_heading('4. Knowledge Base Quality Assessment', level=1)
doc.add_paragraph('Current knowledge base contains 27 documents across 4 categories:')

kb_table = doc.add_table(rows=1, cols=3)
kb_table.style = 'Table Grid'
hdr = kb_table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Count'
hdr[2].text = 'Source'
kb_data = [
    ('Explainers', '4', 'Official Vena Documentation'),
    ('How-To Guides', '10', 'Official Vena Documentation'),
    ('Reference', '5', 'Official Vena Documentation'),
    ('Troubleshooting', '2', 'Official Vena Documentation'),
    ('Custom Guides', '6', 'Internal Documentation'),
]
for cat, count, source in kb_data:
    row = kb_table.add_row().cells
    row[0].text = cat
    row[1].text = count
    row[2].text = source

doc.add_paragraph('')
doc.add_paragraph('Assessment: Knowledge base demonstrates the concept effectively. Production deployment would require expansion to 100+ documents.')

# 5. Recommendations
doc.add_heading('5. Recommendations for Phase 1', level=1)

doc.add_heading('5.1 Technical Recommendations', level=2)
tech_recs = [
    'Deploy to Microsoft Teams for enterprise accessibility',
    'Implement persistent vector storage (Pinecone or Azure Cognitive Search)',
    'Add user authentication via corporate SSO',
    'Implement feedback collection (thumbs up/down)',
    'Add usage analytics and monitoring',
]
for r in tech_recs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.2 Content Recommendations', level=2)
content_recs = [
    'Expand knowledge base to 100+ documents',
    'Include VenaQL code pattern library',
    'Add historical issue resolutions from support tickets',
    'Establish document update and curation process',
]
for r in content_recs:
    doc.add_paragraph(r, style='List Bullet')

# 6. Risk Assessment
doc.add_heading('6. Risk Assessment', level=1)

risk_table = doc.add_table(rows=1, cols=4)
risk_table.style = 'Table Grid'
hdr = risk_table.rows[0].cells
hdr[0].text = 'Risk'
hdr[1].text = 'Likelihood'
hdr[2].text = 'Impact'
hdr[3].text = 'Mitigation'
risks = [
    ('Incorrect answers', 'Medium', 'High', 'Source citations enable verification'),
    ('API cost overruns', 'Low', 'Medium', 'Usage monitoring and rate limiting'),
    ('Outdated knowledge base', 'Medium', 'Medium', 'Regular content review process'),
    ('User adoption', 'Low', 'Medium', 'Phased rollout with champions'),
]
for risk, likelihood, impact, mitigation in risks:
    row = risk_table.add_row().cells
    row[0].text = risk
    row[1].text = likelihood
    row[2].text = impact
    row[3].text = mitigation

# 7. Go/No-Go
doc.add_heading('7. Go/No-Go Recommendation', level=1)
doc.add_paragraph('RECOMMENDATION: GO')
doc.add_paragraph('')
doc.add_paragraph('Rationale:')
rationale = [
    'POC successfully demonstrates core value proposition',
    'Technical feasibility confirmed with working deployment',
    'Semantic search provides significant improvement over traditional documentation',
    'Clear path to production deployment identified',
    'Estimated ROI positive within 6 months',
    'Low implementation risk with established technology stack',
]
for r in rationale:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('7.1 Next Steps (if GO)', level=2)
steps = [
    'Approve Phase 1 budget (estimated GBP 8,000-32,000 annual)',
    'Assign resources (1 ML Engineer, 1 Vena SME, 0.5 Technical Writer)',
    'Begin knowledge base expansion',
    'Initiate Microsoft Teams integration planning',
    'Define success metrics and KPIs',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# Appendix
doc.add_heading('Appendix: Live Demo', level=1)
doc.add_paragraph('POC Application URL: https://4naj96qswvzvmncgpved8x.streamlit.app')
doc.add_paragraph('GitHub Repository: https://github.com/Milesy1/vena-rag-bot')
doc.add_paragraph('')
doc.add_paragraph('Note: Click Build Knowledge Base after app restart.')

doc.save(r'C:\Users\Miles\Desktop\Projects\Vena\Documents\POC_Findings_Report.docx')
print('Done! Report saved.')

